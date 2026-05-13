#!/usr/bin/env python3
"""Quick extraction of Gemini docx"""
import sys
import subprocess
from pathlib import Path

# Install python-docx if needed
try:
    from docx import Document
except:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

gemini_file = Path(r"c:\_ReverseEngineering\spColumnReversed-Solution\Gemini_NvsM_Test_Package_v2.docx")

if not gemini_file.exists():
    print(f"ERROR: {gemini_file} not found")
    sys.exit(1)

try:
    doc = Document(str(gemini_file))
    
    print("=" * 100)
    print("GEMINI DOCUMENT CONTENT")
    print("=" * 100)
    
    # Extract all paragraphs
    print(f"\n📄 Total paragraphs: {len(doc.paragraphs)}")
    print(f"📊 Total tables: {len(doc.tables)}")
    
    print("\nPARAGRAPH CONTENT (first 100 lines):")
    print("-" * 100)
    for i, para in enumerate(doc.paragraphs[:100]):
        if para.text.strip():
            print(f"{i+1:4d}. {para.text[:100]}")
    
    print("\n\nTABLE CONTENT:")
    print("-" * 100)
    for table_idx, table in enumerate(doc.tables[:10]):
        print(f"\nTABLE {table_idx}: ({len(table.rows)} rows × {len(table.columns)} cols)")
        print("Header row:")
        for cell in table.rows[0].cells:
            print(f"  {cell.text[:30]}", end=" ")
        print("\nFirst 5 data rows:")
        for row_idx, row in enumerate(table.rows[1:6]):
            row_text = " | ".join([cell.text[:15] for cell in row.cells])
            print(f"  {row_idx}: {row_text}")
    
    print("\n✓ Document read successfully")
    
except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
