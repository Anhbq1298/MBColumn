#!/usr/bin/env python3
"""
Minimal Gemini docx extraction - just the essentials
Saves results to a file that we can read
"""
import sys
import os
import json

gemini_file = r"c:\_ReverseEngineering\spColumnReversed-Solution\Gemini_NvsM_Test_Package_v2.docx"
output_file = r"c:\_ReverseEngineering\spColumnReversed-Solution\gemini_extracted.json"

print(f"Checking file: {gemini_file}")

if not os.path.exists(gemini_file):
    print(f"ERROR: File not found")
    sys.exit(1)

print(f"✓ File found, size: {os.path.getsize(gemini_file)} bytes")

try:
    from docx import Document
    print("✓ python-docx available")
    
    doc = Document(gemini_file)
    print(f"✓ Opened successfully")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")
    
    # Extract key information
    data = {
        "paragraphs": [],
        "tables": []
    }
    
    # Get first 150 paragraphs
    for i, para in enumerate(doc.paragraphs[:150]):
        text = para.text.strip()
        if text:
            data["paragraphs"].append({
                "index": i,
                "text": text
            })
    
    # Get first 10 tables
    for t_idx, table in enumerate(doc.tables[:10]):
        table_data = {
            "index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "content": []
        }
        
        # Get first 20 rows
        for r_idx, row in enumerate(table.rows[:20]):
            row_cells = [cell.text.strip() for cell in row.cells]
            table_data["content"].append(row_cells)
        
        data["tables"].append(table_data)
    
    # Save to JSON
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    print(f"\n✓ Extracted data saved to: {output_file}")
    print(f"  Paragraphs extracted: {len(data['paragraphs'])}")
    print(f"  Tables extracted: {len(data['tables'])}")
    
    # Print first 50 paragraphs to stdout
    print("\n" + "="*100)
    print("FIRST 50 PARAGRAPHS:")
    print("="*100)
    for item in data["paragraphs"][:50]:
        print(f"{item['index']:3d}. {item['text'][:95]}")

except ImportError:
    print("ERROR: python-docx not installed")
    print("Installing...")
    os.system("pip install python-docx -q")
    print("Please run this script again")
except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
