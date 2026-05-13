#!/usr/bin/env python3
"""
Extract and analyze Gemini_NvsM_Test_Package_v2.docx
Identify section geometry, materials, rebar config, and reference PMM data.
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document


def extract_gemini_document(docx_path):
    """Extract all content from Gemini document"""
    print(f"Opening: {docx_path}\n")
    doc = Document(docx_path)
    
    print("=" * 80)
    print("DOCUMENT STRUCTURE ANALYSIS")
    print("=" * 80)
    
    # Extract all text
    print(f"\n📄 PARAGRAPHS ({len(doc.paragraphs)} total):")
    print("-" * 80)
    
    gemini_config = {
        'section_geometry': {},
        'materials': {},
        'rebar': {},
        'design_code': None,
        'theta_data': {}
    }
    
    for i, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
        text = para.text.strip()
        if text:
            print(f"{i:3d}. {text[:100]}")
            
            # Extract metadata
            if 'section' in text.lower() or 'dimensions' in text.lower():
                gemini_config['section_geometry'][f'line_{i}'] = text
            if 'material' in text.lower() or "f'" in text or 'fy' in text:
                gemini_config['materials'][f'line_{i}'] = text
            if 'rebar' in text.lower() or 'bar' in text.lower() or 'reinforcement' in text.lower():
                gemini_config['rebar'][f'line_{i}'] = text
            if 'ACI' in text or 'Eurocode' in text or 'EC2' in text:
                gemini_config['design_code'] = text
    
    # Extract tables
    print(f"\n📊 TABLES ({len(doc.tables)} total):")
    print("-" * 80)
    
    for table_idx, table in enumerate(doc.tables[:5]):  # First 5 tables
        print(f"\nTable {table_idx}:")
        print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        
        # Print first 5 rows
        for row_idx, row in enumerate(table.rows[:5]):
            cells = [cell.text.strip()[:20] for cell in row.cells]
            print(f"    Row {row_idx}: {cells}")
    
    return gemini_config


def main():
    repo_root = Path(__file__).parent
    gemini_file = repo_root / "Gemini_NvsM_Test_Package_v2.docx"
    
    if not gemini_file.exists():
        print(f"ERROR: {gemini_file} not found")
        sys.exit(1)
    
    config = extract_gemini_document(str(gemini_file))
    
    print("\n" + "=" * 80)
    print("EXTRACTED CONFIGURATION")
    print("=" * 80)
    print(f"\nSection Geometry: {config['section_geometry']}")
    print(f"Materials: {config['materials']}")
    print(f"Rebar: {config['rebar']}")
    print(f"Design Code: {config['design_code']}")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
