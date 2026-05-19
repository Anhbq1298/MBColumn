import docx
import re
import json

def extract_target_angles(docx_path, target_angles):
    doc = docx.Document(docx_path)
    all_data = {}
    current_angle = None
    
    # First, look for paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.search(r'Theta\s*\(Degrees\)\s*(\d+)', text, re.IGNORECASE)
        if match:
            current_angle = int(match.group(1))
            if current_angle in target_angles:
                all_data[current_angle] = []
            else:
                current_angle = None
            continue
            
        if current_angle is not None:
            # Match data lines like "-104.7713	 573.5886"
            nums = re.findall(r'[-+]?\d*\.\d+|\b\d+\b', text)
            if len(nums) >= 2:
                try:
                    p = float(nums[0])
                    m = float(nums[1])
                    all_data[current_angle].append({"P": p, "M": m})
                except ValueError:
                    pass

    # Some tables might contain the data
    for table in doc.tables:
        # Check if the text before the table contains the angle
        # But for this doc, we saw the angle in paragraphs.
        pass

    return all_data

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\ref\Gemini_NvsM_Test_Package_v2.docx"
    targets = [0, 16, 75, 90, 135]
    data = extract_target_angles(path, targets)
    
    # Save to JSON for C# to read
    with open(r"C:\Users\NCPC\OneDrive - Meinhardt Singapore Pte Ltd\_ReverseEngineering\refReversed-Solution\tests\MBColumn.Tests\benchmark_multi_angle.json", "w") as f:
        json.dump(data, f, indent=2)
        
    for angle in targets:
        if angle in data:
            print(f"Angle {angle}: {len(data[angle])} points extracted.")
        else:
            print(f"Angle {angle}: NOT FOUND")
