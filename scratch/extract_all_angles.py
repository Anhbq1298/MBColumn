import docx
import re

def extract_benchmark_data(docx_path):
    doc = docx.Document(docx_path)
    data = {}
    current_angle = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        # Look for headers like "Angle = 0", "Angle: 16", etc.
        match = re.search(r'(?:Angle|Theta)\s*[=:]\s*(\d+)', text, re.IGNORECASE)
        if match:
            current_angle = int(match.group(1))
            data[current_angle] = []
            continue
            
        # If we have an angle, look for lines with two numbers (P and M)
        if current_angle is not None:
            # Match lines like "-7648.306 649.0892" or "0 753.5"
            # Using regex to find pairs of floats
            nums = re.findall(r'[-+]?\d*\.\d+|\b\d+\b', text)
            if len(nums) >= 2:
                try:
                    p = float(nums[0])
                    m = float(nums[1])
                    data[current_angle].append((p, m))
                except ValueError:
                    pass
                    
    # Also check tables
    for table in doc.tables:
        # Tables often have a header row then data
        # We need to find which angle this table belongs to
        # Usually the angle is in the text BEFORE the table
        # But let's just dump all table data to see what's there
        pass

    return data

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\ref\Gemini_NvsM_Test_Package_v2.docx"
    all_data = extract_benchmark_data(path)
    for angle in sorted(all_data.keys()):
        print(f"--- Angle {angle} ---")
        for p, m in all_data[angle]:
            print(f"{p:>12} {m:>12}")
