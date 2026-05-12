import docx

def dump_docx(docx_path):
    doc = docx.Document(docx_path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{i}: {para.text}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTABLE {i}:")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\s-conc\Gemini_NvsM_Test_Package_v2.docx"
    dump_docx(path)
