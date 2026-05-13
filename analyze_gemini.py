#!/usr/bin/env python3
"""
Comprehensive analysis of Gemini_NvsM_Test_Package_v2.docx
Extracts: section geometry, materials, rebar, design code, and PMM reference data
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document


def extract_all_text(docx_path):
    """Extract all text from document"""
    doc = Document(docx_path)
    text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
    
    return text


def extract_from_tables(docx_path):
    """Extract data from tables"""
    doc = Document(docx_path)
    tables_data = []
    
    for table_idx, table in enumerate(doc.tables):
        table_text = []
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            table_text.append(row_data)
        tables_data.append({
            'index': table_idx,
            'rows': table_text,
            'row_count': len(table.rows),
            'col_count': len(table.columns)
        })
    
    return tables_data


def parse_parameters(all_text):
    """Extract test parameters from text"""
    params = {
        'section_width': None,
        'section_height': None,
        'rebar_count': None,
        'rebar_diameter': None,
        'fc_mpa': None,
        'fy_mpa': None,
        'design_code': None,
        'unit_system': None
    }
    
    full_text = ' '.join(all_text).lower()
    
    # Look for dimensions (mm)
    dim_patterns = [
        r'(\d+)\s*x\s*(\d+)\s*mm',
        r'(\d+)\s*×\s*(\d+)\s*mm',
        r'section\s*=\s*(\d+)\s*x\s*(\d+)',
        r'(\d+)\s*mm\s*x\s*(\d+)\s*mm',
    ]
    
    for pattern in dim_patterns:
        matches = re.findall(pattern, full_text)
        if matches:
            params['section_width'] = float(matches[0][0])
            params['section_height'] = float(matches[0][1])
            break
    
    # Look for rebar
    rebar_patterns = [
        r'(\d+)\s*(?:x|×)\s*(\d+)\s*mm\s*(?:bars?|rebar)',
        r'(\d+)\s*bars?\s*(?:of|with)\s*(\d+)\s*mm',
        r'(\d+)\s*(\d+)mm?\s*rebar',
        r'rebar:\s*(\d+)\s*(\d+)',
    ]
    
    for pattern in rebar_patterns:
        matches = re.findall(pattern, full_text)
        if matches:
            params['rebar_count'] = int(matches[0][0])
            params['rebar_diameter'] = float(matches[0][1])
            break
    
    # Look for materials
    if "f'c" in full_text or 'fc' in full_text or 'fck' in full_text:
        fc_match = re.search(r"f'?c\s*=?\s*(\d+)\s*(?:mpa|mpa|n/mm)", full_text)
        if fc_match:
            params['fc_mpa'] = float(fc_match.group(1))
    
    if 'fy' in full_text or 'fyk' in full_text:
        fy_match = re.search(r"fy\s*=?\s*(\d+)\s*(?:mpa|mpa|n/mm)", full_text)
        if fy_match:
            params['fy_mpa'] = float(fy_match.group(1))
    
    # Look for design code
    if 'aci' in full_text or 'aci318' in full_text:
        params['design_code'] = 'ACI 318'
    elif 'eurocode' in full_text or 'ec2' in full_text:
        params['design_code'] = 'Eurocode 2'
    elif 'as' in full_text or 'as3600' in full_text:
        params['design_code'] = 'AS 3600'
    
    return params


def analyze_pmm_data(all_text, tables_data):
    """Extract PMM curve data"""
    pmm_data = {}
    
    # Look for theta sections
    for text_item in all_text:
        # Match patterns like "Theta = 0", "Θ = 0°", etc.
        theta_match = re.search(r'[Θθ](?:eta)?\s*(?:\(.*?\))?\s*=\s*(\d+)', text_item)
        if theta_match:
            theta = int(theta_match.group(1))
            pmm_data[theta] = []
    
    # Extract from tables if they look like PMM data
    for table in tables_data:
        # Try to identify PMM table (typically has P, M or Px, Mx, My columns)
        if table['row_count'] > 2 and table['col_count'] >= 2:
            rows = table['rows']
            
            # Check header for P, M, N, Moment
            header = rows[0] if rows else []
            header_text = ' '.join(header).lower()
            
            is_pmm_table = any(keyword in header_text for keyword in ['p', 'm', 'n', 'moment', 'axial', 'load'])
            
            if is_pmm_table:
                print(f"\n📊 Found PMM Table (Table {table['index']}):")
                print(f"   Rows: {table['row_count']}, Cols: {table['col_count']}")
                print(f"   Header: {header}")
    
    return pmm_data


def main():
    gemini_file = Path(__file__).parent / "Gemini_NvsM_Test_Package_v2.docx"
    
    if not gemini_file.exists():
        print(f"ERROR: {gemini_file} not found")
        return 1
    
    print("=" * 100)
    print("GEMINI TEST PACKAGE ANALYSIS")
    print("=" * 100)
    
    print(f"\n📄 Reading: {gemini_file}")
    
    # Extract content
    all_text = extract_all_text(str(gemini_file))
    tables_data = extract_from_tables(str(gemini_file))
    
    print(f"✓ Document contains {len(all_text)} text paragraphs")
    print(f"✓ Document contains {len(tables_data)} tables")
    
    # Display first 30 lines
    print("\n📋 FIRST 30 LINES OF DOCUMENT:")
    print("-" * 100)
    for i, line in enumerate(all_text[:30]):
        if line:
            print(f"{i:3d}. {line[:95]}")
    
    # Parse parameters
    print("\n" + "=" * 100)
    print("EXTRACTED TEST PARAMETERS")
    print("=" * 100)
    
    params = parse_parameters(all_text)
    print(f"\nSection Geometry:")
    print(f"  Width: {params['section_width']} mm")
    print(f"  Height: {params['section_height']} mm")
    
    print(f"\nReinforcement:")
    print(f"  Count: {params['rebar_count']} bars")
    print(f"  Diameter: {params['rebar_diameter']} mm")
    
    print(f"\nMaterial Properties:")
    print(f"  f'c: {params['fc_mpa']} MPa")
    print(f"  fy: {params['fy_mpa']} MPa")
    
    print(f"\nDesign Code: {params['design_code']}")
    
    # Analyze tables
    print("\n" + "=" * 100)
    print("TABLE ANALYSIS")
    print("=" * 100)
    
    for table in tables_data:
        print(f"\nTable {table['index']}: {table['row_count']} rows × {table['col_count']} cols")
        if table['rows']:
            print(f"  Header: {table['rows'][0]}")
            if len(table['rows']) > 1:
                print(f"  Sample row: {table['rows'][1]}")
    
    # Try to find PMM data
    print("\n" + "=" * 100)
    print("PMM DATA EXTRACTION")
    print("=" * 100)
    
    pmm_data = analyze_pmm_data(all_text, tables_data)
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
