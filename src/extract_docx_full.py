import docx
import json

def extract_all(docx_path):
    doc = docx.Document(docx_path)
    print(f"Number of tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row in table.rows[:15]: # Show first 15 rows
            print([cell.text.strip().replace('\n', ' ') for cell in row.cells])
        if len(table.rows) > 15:
            print(f"... and {len(table.rows)-15} more rows")

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\s-conc\Gemini_NvsM_Test_Package_v2.docx"
    extract_all(path)
