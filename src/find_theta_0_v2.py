import docx

def find_theta_0(docx_path):
    doc = docx.Document(docx_path)
    for idx, p in enumerate(doc.paragraphs):
        if 'Theta (Degrees)' in p.text and ('\t 0' in p.text or ' 0' in p.text) and '106' not in p.text:
            print(f"Theta 0 found at index {idx}")
            for i in range(idx, min(idx + 105, len(doc.paragraphs))):
                print(doc.paragraphs[i].text)
            return

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\s-conc\Gemini_NvsM_Test_Package_v2.docx"
    find_theta_0(path)
