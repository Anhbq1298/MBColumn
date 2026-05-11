import docx

def find_theta_0(docx_path):
    doc = docx.Document(docx_path)
    found = False
    for p in doc.paragraphs:
        if 'Theta (Degrees)\t 0' in p.text or 'Theta (Degrees) 0' in p.text:
            found = True
            print(f"Theta 0 found in paragraph!")
            # Print the next few paragraphs
            idx = doc.paragraphs.index(p)
            for i in range(idx, idx + 110):
                if i < len(doc.paragraphs):
                    print(doc.paragraphs[i].text)
            break
    if not found:
        print("Theta 0 not found")

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\s-conc\Gemini_NvsM_Test_Package_v2.docx"
    find_theta_0(path)
