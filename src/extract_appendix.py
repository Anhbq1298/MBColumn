import docx

def extract_start_of_appendix(docx_path):
    doc = docx.Document(docx_path)
    found = False
    for idx, p in enumerate(doc.paragraphs):
        if 'Appendix A' in p.text:
            found = True
            print(f"Appendix A found at index {idx}")
            for i in range(idx, min(idx + 100, len(doc.paragraphs))):
                print(doc.paragraphs[i].text)
            return

if __name__ == "__main__":
    path = r"C:\Users\NCPC\Desktop\_ReverseEngineering\s-conc\Gemini_NvsM_Test_Package_v2.docx"
    extract_start_of_appendix(path)
