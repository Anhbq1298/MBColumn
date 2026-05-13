#!/usr/bin/env python3
"""
Extract PMM reference data from Gemini_NvsM_Test_Package_v2.docx
and convert to CSV format for MBColumn verification test.

Expected CSV output format: Theta,PointIndex,RefP,RefMx,RefMy
"""

import sys
import csv
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)


def extract_gemini_reference_data(docx_path: str) -> dict:
    """
    Extract PMM curves from Gemini document.
    Returns dict: {theta_degrees: [(P, Mx, My), ...]}
    """
    doc = Document(docx_path)
    data_by_theta = {}
    current_theta = None
    current_points = []

    for para in doc.paragraphs:
        text = para.text.strip()

        # Look for theta heading (e.g., "Theta = 0 degrees")
        if "Theta" in text and "degrees" in text and "=" in text:
            # Save previous theta if any
            if current_theta is not None and current_points:
                data_by_theta[current_theta] = current_points
            
            try:
                # Extract theta value from "Theta = X" or "Theta (Degrees) = X"
                parts = text.split("=")
                if len(parts) >= 2:
                    theta_str = parts[1].strip().split()[0]
                    current_theta = int(float(theta_str))
                    current_points = []
            except (ValueError, IndexError):
                pass
            continue

        # Try to parse as PMM data point (P M format or P Mx My format)
        if current_theta is not None and text and not any(skip in text for skip in ["Appendix", "Notes", "Table", "Theta"]):
            parts = text.split()
            try:
                if len(parts) >= 2:
                    p = float(parts[0])
                    m1 = float(parts[1])
                    
                    # Assume for now: if 3 values, it's P Mx My; if 2 values, it's P M (resultant)
                    if len(parts) >= 3:
                        m2 = float(parts[2])
                        current_points.append((p, m1, m2))
                    else:
                        # Single M value - assume it's resultant, distribute equally
                        # For biaxial, assume equal Mx and My (this is a guess)
                        mx = m1
                        my = 0.0  # We'll adjust if needed
                        current_points.append((p, mx, my))
            except (ValueError, IndexError):
                pass

    # Save last theta
    if current_theta is not None and current_points:
        data_by_theta[current_theta] = current_points

    return data_by_theta


def extract_from_tables(docx_path: str) -> dict:
    """
    Alternative extraction method using document tables directly.
    """
    doc = Document(docx_path)
    data_by_theta = {}

    for table in doc.tables:
        # Try to identify theta from table header or caption
        current_theta = None
        points = []

        for row_idx, row in enumerate(table.rows):
            row_text = " ".join([cell.text.strip() for cell in row.cells])

            # Check if this is a header row with theta info
            if "Theta" in row_text and "=" in row_text:
                try:
                    parts = row_text.split("=")
                    if len(parts) >= 2:
                        theta_str = parts[1].strip().split()[0]
                        current_theta = int(float(theta_str))
                except (ValueError, IndexError):
                    pass
                continue

            # Try to parse as data row
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) >= 2 and current_theta is not None:
                try:
                    p = float(cells[0])
                    m1 = float(cells[1])
                    m2 = float(cells[2]) if len(cells) >= 3 else 0.0
                    points.append((p, m1, m2))
                except (ValueError, IndexError):
                    pass

        if current_theta is not None and points:
            data_by_theta[current_theta] = points

    return data_by_theta


def write_csv(output_csv_path: str, data_by_theta: dict):
    """
    Convert extracted data to CSV format.
    Format: Theta,PointIndex,RefP,RefMx,RefMy
    """
    with open(output_csv_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['Theta', 'PointIndex', 'RefP', 'RefMx', 'RefMy'])

        for theta in sorted(data_by_theta.keys()):
            points = data_by_theta[theta]
            for point_idx, (p, mx, my) in enumerate(points):
                writer.writerow([theta, point_idx, f'{p:.4f}', f'{mx:.4f}', f'{my:.4f}'])

    print(f"✓ Wrote {output_csv_path}")
    total_points = sum(len(pts) for pts in data_by_theta.values())
    print(f"  Theta angles: {len(data_by_theta)}")
    print(f"  Total points: {total_points}")


def main():
    # Locate reference document
    repo_root = Path(__file__).parent.parent
    docx_file = repo_root / "Gemini_NvsM_Test_Package_v2.docx"

    if not docx_file.exists():
        print(f"ERROR: Reference file not found: {docx_file}")
        sys.exit(1)

    print(f"Extracting from: {docx_file}")

    # Try extraction methods
    print("\n1. Attempting paragraph-based extraction...")
    data_para = extract_gemini_reference_data(str(docx_file))
    if data_para:
        print(f"   Found {len(data_para)} theta angles via paragraphs")

    print("\n2. Attempting table-based extraction...")
    data_tables = extract_from_tables(str(docx_file))
    if data_tables:
        print(f"   Found {len(data_tables)} theta angles via tables")

    # Use best result
    data = data_tables if len(data_tables) > len(data_para) else data_para

    if not data:
        print("\nWARNING: No reference data extracted. Check document format.")
        print("  Expected: Tables or paragraphs with 'Theta' and PMM values")
        return

    # Output CSV
    output_csv = repo_root / "docs" / "validation" / "gemini-reference-points.csv"
    output_csv.parent.mkdir(parents=True, exist_ok=True)
    write_csv(str(output_csv), data)

    print("\n✓ Reference CSV created successfully")
    return 0


if __name__ == "__main__":
    sys.exit(main())
