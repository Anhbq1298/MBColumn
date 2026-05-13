#!/usr/bin/env python3
"""
Very simple read of Gemini docx to dump its structure
"""
import sys
import os

gemini_file = r"c:\_ReverseEngineering\spColumnReversed-Solution\Gemini_NvsM_Test_Package_v2.docx"

print(f"Checking if file exists: {gemini_file}")
if os.path.exists(gemini_file):
    print(f"✓ File found, size: {os.path.getsize(gemini_file)} bytes")
    
    # Try to read with python-docx
    try:
        from docx import Document
        doc = Document(gemini_file)
        
        print(f"\n✓ Successfully opened with python-docx")
        print(f"  Paragraphs: {len(doc.paragraphs)}")
        print(f"  Tables: {len(doc.tables)}")
        
        # Dump first 100 paragraphs
        print("\nFIRST 100 PARAGRAPHS:")
        print("=" * 100)
        for i, para in enumerate(doc.paragraphs[:100]):
            text = para.text.strip()
            if text:
                print(f"{i+1:4d}. {text}")
        
        # Dump first 5 tables
        print("\n\nFIRST 5 TABLES:")
        print("=" * 100)
        for t_idx, table in enumerate(doc.tables[:5]):
            print(f"\nTABLE {t_idx}: {len(table.rows)} rows × {len(table.columns)} cols")
            for r_idx, row in enumerate(table.rows[:10]):
                cells = " | ".join([c.text[:20] for c in row.cells])
                print(f"  Row {r_idx}: {cells}")
                
    except ImportError:
        print("python-docx not installed, installing...")
        os.system("pip install python-docx -q")
        print("Please run this script again")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
else:
    print(f"✗ File not found")
    sys.exit(1)
