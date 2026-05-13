#!/usr/bin/env python3
"""
Extract test parameters and PMM data from Gemini_NvsM_Test_Package_v2.docx
"""

import sys
import os
from pathlib import Path
import csv
import re

# Try to import python-docx
try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

def extract_docx_content(file_path):
    """Extract all text and tables from docx file"""
    doc = Document(file_path)
    
    full_text = []
    tables_data = []
    
    # Extract paragraphs
    print("\n" + "="*80)
    print("DOCUMENT TEXT CONTENT")
    print("="*80)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            full_text.append(para.text)
    
    # Extract tables
    print("\n" + "="*80)
    print(f"FOUND {len(doc.tables)} TABLE(S)")
    print("="*80)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {table_idx + 1} ---")
        table_rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_data)
            print(row_data)
        tables_data.append(table_rows)
    
    return full_text, tables_data

def parse_test_parameters(full_text):
    """Parse test parameters from extracted text"""
    params = {
        'section_dimensions': {},
        'rebar_config': {},
        'material_properties': {},
        'design_code': None,
        'loads': {},
        'units': None,
        'raw_text': full_text
    }
    
    text_combined = ' '.join(full_text).lower()
    
    # Look for common keywords
    if 'aci' in text_combined:
        params['design_code'] = 'ACI'
    elif 'eurocode' in text_combined or 'ec2' in text_combined:
        params['design_code'] = 'Eurocode'
    
    if 'mm' in text_combined:
        params['units'] = 'mm'
    elif 'inch' in text_combined or '"' in text_combined:
        params['units'] = 'inches'
    
    return params

def main():
    docx_path = r"c:\_ReverseEngineering\spColumnReversed-Solution\Gemini_NvsM_Test_Package_v2.docx"
    
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}")
        return
    
    print(f"Reading: {docx_path}")
    full_text, tables_data = extract_docx_content(docx_path)
    
    print("\n" + "="*80)
    print("PARSED PARAMETERS")
    print("="*80)
    params = parse_test_parameters(full_text)
    
    print(f"Design Code: {params['design_code']}")
    print(f"Units: {params['units']}")
    
    # Save tables to CSV if they contain PMM data
    if tables_data:
        print("\n" + "="*80)
        print("SAVING TABLE DATA TO CSV")
        print("="*80)
        
        for table_idx, table_rows in enumerate(tables_data):
            if len(table_rows) > 1:
                csv_filename = f"gemini_table_{table_idx}.csv"
                csv_path = os.path.join(
                    r"c:\_ReverseEngineering\spColumnReversed-Solution\docs\validation",
                    csv_filename
                )
                
                # Create directory if needed
                os.makedirs(os.path.dirname(csv_path), exist_ok=True)
                
                with open(csv_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerows(table_rows)
                
                print(f"Saved: {csv_path}")
                print(f"  Rows: {len(table_rows)}, Columns: {len(table_rows[0]) if table_rows else 0}")

if __name__ == '__main__':
    main()
